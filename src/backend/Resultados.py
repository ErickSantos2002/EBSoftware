import os
import sys
import pandas as pd
import tempfile
import configparser
from docx import Document
from win32com.client import Dispatch
from datetime import datetime
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfgen import canvas
from src.backend.db import conectar

if getattr(sys, 'frozen', False):
    # Diretório do executável (PyInstaller)
    BASE_DIR = sys._MEIPASS
else:
    # Diretório raiz do projeto (quando executado como script)
    BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
RESOURCES_DIR = os.path.join(BASE_DIR, "resources")  # Caminho do diretório resources
INFO_FILE = os.path.join(RESOURCES_DIR, "info.ini")

# Caminho do arquivo específico
ARQUIVO_RESULTADOS = os.path.join(RESOURCES_DIR, "Resultados.csv")

def gerar_laudo(data_emissao, nome, matricula, setor, data_teste, resultado, unidade, limite_baixo, limite_alto, numero_testes, save_path):
    """Gera um laudo técnico em PDF com base no arquivo temporário."""
    try:
        # Caminho do modelo de laudo
        modelo_docx = os.path.join(RESOURCES_DIR, "Modelo Laudo.docx")
        if not os.path.exists(modelo_docx):
            raise FileNotFoundError(f"Modelo de Laudo não encontrado em: {modelo_docx}")

        # Criar uma cópia temporária do modelo
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
            temp_docx_path = temp_docx.name
        doc = Document(modelo_docx)

        # Substituir os placeholders no arquivo temporário
        placeholders = {
            "{{data_hora_emissao}}": data_emissao,
            "{{nome_usuario}}": nome,
            "{{matricula_usuario}}": matricula,
            "{{setor_usuario}}": setor,
            "{{data_hora_teste}}": data_teste,
            "{{resultado_teste}}": resultado,
            "{{unidade}}": unidade,
            "{{limite_baixo}}": limite_baixo,
            "{{limite_alto}}": limite_alto,
            "{{numero_testes}}": numero_testes,
        }

        for p in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, value)

        # Substituir placeholders em tabelas (se houver)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in placeholders.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)

        # Salvar o documento temporário com as alterações
        doc.save(temp_docx_path)

        # Converter o Word temporário para PDF mantendo o layout
        word = Dispatch("Word.Application")
        doc = word.Documents.Open(temp_docx_path)
        doc.SaveAs(save_path, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()

        # Remover o arquivo temporário
        os.remove(temp_docx_path)

        return save_path
    except Exception as e:
        raise Exception(f"Erro ao gerar o laudo: {e}")

    
def carregar_resultados():
    """Carrega resultados do banco de dados."""
    with conectar() as conn:
        cursor = conn.cursor()
        cursor.execute("""
        SELECT id, id_usuario, nome, matricula, setor, data_hora, quantidade_alcool, status
        FROM resultados
        """)
        resultados = [
            {
                "ID do teste": row[0],
                "ID do usuário": row[1],
                "Nome": row[2],
                "Matrícula": row[3],
                "Setor": row[4],
                "Data e hora": row[5],
                "Quantidade de Álcool": row[6],
                "Status": row[7],
            } for row in cursor.fetchall()
        ]
    return resultados


def filtrar_resultados(resultados, periodo=None, usuario=None, status=None):
    """Filtra os resultados com base nos parâmetros fornecidos."""
    filtrados = []

    for resultado in resultados:
        # Filtro por período
        if periodo:
            data_teste = datetime.strptime(resultado["Data e hora"], "%Y/%m/%d %H:%M:%S").date()
            if not (periodo[0] <= data_teste <= periodo[1]):
                continue

        # Filtro por usuário
        if usuario and resultado["Nome"] != usuario:
            continue

        # Filtro por status
        if status and resultado["Status"] != status:
            continue

        filtrados.append(resultado)

    return filtrados

def salvar_em_excel(resultados, caminho_arquivo):
    """Salva os resultados em um arquivo Excel."""
    df = pd.DataFrame(resultados)

    # Formata a coluna de Data e Hora, se existir
    if "Data e hora" in df.columns:
        df["Data e hora"] = pd.to_datetime(df["Data e hora"]).dt.strftime("%d/%m/%y %H:%M:%S")

    with pd.ExcelWriter(caminho_arquivo, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Resultados")
        worksheet = writer.sheets["Resultados"]

        # Ajusta a largura das colunas no Excel
        for col_idx, column_cells in enumerate(worksheet.iter_cols()):
            max_length = max(len(str(cell.value)) for cell in column_cells if cell.value)
            col_letter = worksheet.cell(row=1, column=col_idx + 1).column_letter
            worksheet.column_dimensions[col_letter].width = max_length + 2

def salvar_em_pdf(resultados, caminho_arquivo):
    """Salva os resultados em um arquivo PDF."""
    c = canvas.Canvas(caminho_arquivo, pagesize=landscape(letter))
    largura, altura = landscape(letter)

    # Título
    c.setFont("Helvetica-Bold", 16)
    c.drawString(30, altura - 40, "Registros de Resultados EBS010")

    # Cabeçalhos
    colunas = ["ID do teste", "ID do usuário", "Nome", "Matrícula", "Setor", "Data e hora", "Qtd. Álcool", "Status"]
    larguras_colunas = [60, 80, 100, 90, 90, 160, 90, 80]
    c.setFont("Helvetica-Bold", 10)
    y = altura - 70
    x_inicial = 30
    for i, coluna in enumerate(colunas):
        c.drawString(x_inicial, y, coluna)
        x_inicial += larguras_colunas[i]

    # Dados
    c.setFont("Helvetica", 10)
    y -= 20
    for row in resultados:
        x_inicial = 30
        for i, coluna in enumerate(colunas):
            c.drawString(x_inicial, y, str(row[coluna]))
            x_inicial += larguras_colunas[i]
        y -= 20
        if y < 50:
            c.showPage()
            y = altura - 70
            c.setFont("Helvetica-Bold", 10)
            for i, coluna in enumerate(colunas):
                c.drawString(x_inicial, y, coluna)
                x_inicial += larguras_colunas[i]
            c.setFont("Helvetica", 10)
            y -= 20

    c.save()
